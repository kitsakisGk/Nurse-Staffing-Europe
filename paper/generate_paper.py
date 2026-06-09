"""Convert nurse_staffing_europe.md to a formatted Word document."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent.parent
PAPER_MD = BASE / "paper" / "nurse_staffing_europe.md"
FIGURES = BASE / "figures"
OUTPUT = BASE / "paper" / "nurse_staffing_europe.docx"

FIGURE_MAP = {
    "Figure 1": FIGURES / "screenshots" / "Nurse_Staffing_Overview_Map.png",
    "Figure 2": FIGURES / "screenshots" / "Trends_Regions_Full_Plots.png",
    "Figure 3": FIGURES / "screenshots" / "Clusters_Regression_full_plots.png",
}


def set_cell_background(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_table_from_md(doc, rows):
    """rows: list of lists of cell strings (first row = header)"""
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = cell_text.strip()
            if r_idx == 0:
                set_cell_background(cell, "2C3E50")
                run = cell.paragraphs[0].runs
                if run:
                    run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run[0].font.bold = True
    return table


def build_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "Nurse Staffing and Patient Outcomes in Europe:\n"
        "A Panel Data Analysis Across 36 Countries (2000–2023)"
    )
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Author / meta
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Giorgos Kitsakis\n").font.bold = True
    meta_run2 = meta.add_run(
        "Data sources: OECD Health Statistics · WHO Global Health Observatory · Eurostat · World Bank\n"
        "Repository: https://github.com/kitsakisGk/Nurse-Staffing-Europe"
    )
    meta_run2.font.size = Pt(10)
    meta_run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    with open(PAPER_MD, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    table_rows = []
    in_table = False

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip the markdown title (first line, already added)
        if i == 0 and line.startswith("# "):
            i += 1
            continue

        # Skip author/date/source metadata block at top
        if line.startswith("**Author:**") or line.startswith("**Date:**") or \
           line.startswith("**Data sources:**") or line.startswith("**Repository:**"):
            i += 1
            continue

        # Horizontal rule → page break-ish separator
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Headings
        if line.startswith("## "):
            if in_table:
                add_table_from_md(doc, table_rows)
                table_rows = []
                in_table = False
            add_heading(doc, line[3:], level=1)
            i += 1
            continue

        if line.startswith("### "):
            if in_table:
                add_table_from_md(doc, table_rows)
                table_rows = []
                in_table = False
            add_heading(doc, line[4:], level=2)
            i += 1
            continue

        # Figure image line  ![...](path)
        if line.startswith("!["):
            # Extract figure number for lookup
            fig_num = None
            for key in FIGURE_MAP:
                if key in line:
                    fig_num = key
                    break
            if fig_num and FIGURE_MAP[fig_num].exists():
                doc.add_picture(str(FIGURE_MAP[fig_num]), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            # Next line is the caption (*Figure N: ...*) — style it
            if i < len(lines):
                caption_line = lines[i].strip().strip("*")
                p = doc.add_paragraph(caption_line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.italic = True
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                i += 1
            doc.add_paragraph()
            continue

        # Table rows (pipe-separated)
        if line.startswith("|"):
            cells = [c for c in line.split("|") if c.strip() != ""]
            # Skip separator rows like |---|---|
            if all(set(c.strip()) <= set("-: ") for c in cells):
                i += 1
                in_table = True
                continue
            table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                add_table_from_md(doc, table_rows)
                table_rows = []
                in_table = False

        # Italic caption lines following table
        if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
            p = doc.add_paragraph(line.strip("*"))
            for run in p.runs:
                run.font.italic = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            i += 1
            continue

        # Numbered references (lines starting with digit + dot)
        if len(line) > 2 and line[0].isdigit() and line[1] in ".0123456789":
            p = doc.add_paragraph(style="List Number")
            # strip leading number
            text = line.lstrip("0123456789").lstrip(". ")
            add_inline_formatting(p, text)
            i += 1
            continue

        # Bullet list
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatting(p, line[2:])
            i += 1
            continue

        # Numbered list (1. 2. etc.)
        import re
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\. ", "", line)
            add_inline_formatting(p, text)
            i += 1
            continue

        # Block quote (equation)
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line[2:])
            run.font.italic = True
            i += 1
            continue

        # Empty line
        if line.strip() == "":
            i += 1
            continue

        # Normal paragraph
        if line.strip():
            p = doc.add_paragraph()
            add_inline_formatting(p, line.strip())

        i += 1

    if in_table:
        add_table_from_md(doc, table_rows)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


def add_inline_formatting(paragraph, text):
    """Handle **bold** and *italic* inline markdown."""
    import re
    # Pattern: **bold**, *italic*, or plain text
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            paragraph.add_run(part)


if __name__ == "__main__":
    build_docx()
